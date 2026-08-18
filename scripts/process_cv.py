#!/usr/bin/env python3
"""
scripts/process_cv.py
=====================================================================
Reads an uploaded CV (.pdf / .docx / .txt), sends its text plus the
user's enhancement instruction to Google Gemini, and writes an
optimized CV to the output directory.

Invoked by .github/workflows/cv-enhancer.yml:

    python scripts/process_cv.py \
        --input  inbox/<filename> \
        --prompt "<enhancement instruction>" \
        --email  "<requester email>" \
        --outdir output

Environment:
    GEMINI_API_KEY   Gemini API key (mapped from secrets.LLM_API_KEY)

Dependencies:
    google-genai, python-docx, pypdf
=====================================================================
"""

import argparse
import os
import sys
from datetime import datetime, timezone
from pathlib import Path


# ---------------------------------------------------------------------------
# CV text extraction (.txt / .docx / .pdf)
# ---------------------------------------------------------------------------
def extract_text(path: Path) -> str:
    """Return plain text from a .txt, .docx or .pdf file."""
    ext = path.suffix.lower()

    if ext == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")

    if ext == ".docx":
        try:
            from docx import Document  # python-docx
        except ImportError:
            sys.exit("::error::python-docx not installed; cannot read .docx")
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        # include table cell text as well
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(t for t in parts if t and t.strip())

    if ext == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            sys.exit("::error::pypdf not installed; cannot read .pdf")
        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "") for page in reader.pages)

    sys.exit(f"::error::Unsupported file type: {ext}")


# ---------------------------------------------------------------------------
# Gemini call
# ---------------------------------------------------------------------------
def build_prompt(cv_text: str, user_instruction: str) -> str:
    """Compose the instruction sent to Gemini."""
    return f"""You are an expert technical resume writer and career coach.

TASK
Rewrite and enhance the CV below according to the user's instruction.
Preserve every real fact (employers, dates, education, certifications) —
never invent experience. Improve clarity, impact and ATS-friendliness:
  - Lead bullets with strong action verbs and quantified outcomes.
  - Tighten wording; remove fluff and redundancy.
  - Group skills logically and surface the most relevant ones first.
  - Keep a clean, professional structure with clear section headings.

USER INSTRUCTION
{user_instruction}

RETURN FORMAT
Return the full enhanced CV in clean Markdown. After the CV, add a short
"## What Changed" section summarising the key improvements you made.

--- ORIGINAL CV START ---
{cv_text}
--- ORIGINAL CV END ---
"""


def enhance_with_gemini(cv_text: str, user_instruction: str) -> str:
    """Send the CV to Gemini and return the enhanced markdown."""
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        sys.exit("::error::GEMINI_API_KEY is not set (map secrets.LLM_API_KEY).")

    try:
        # Official Google Gen AI SDK
        from google import genai
        from google.genai import types
    except ImportError:
        sys.exit("::error::google-genai not installed. Run: pip install google-genai")

    client = genai.Client(api_key=api_key)
    model = os.environ.get("GEMINI_MODEL", "gemini-2.0-flash")

    response = client.models.generate_content(
        model=model,
        contents=build_prompt(cv_text, user_instruction),
        config=types.GenerateContentConfig(
            temperature=0.4,
            max_output_tokens=4096,
        ),
    )

    text = getattr(response, "text", None)
    if not text:
        sys.exit("::error::Gemini returned an empty response.")
    return text


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main() -> None:
    parser = argparse.ArgumentParser(description="Gemini-powered CV enhancer")
    parser.add_argument("--input", required=True, help="Path to uploaded CV")
    parser.add_argument("--prompt", required=True, help="Enhancement instruction")
    parser.add_argument("--email", default="", help="Requester email (for reference)")
    parser.add_argument("--outdir", default="output", help="Output directory")
    args = parser.parse_args()

    in_path = Path(args.input)
    if not in_path.exists():
        sys.exit(f"::error::Input file not found: {in_path}")

    out_dir = Path(args.outdir)
    out_dir.mkdir(parents=True, exist_ok=True)

    print(f"[1/3] Extracting text from {in_path.name} …")
    cv_text = extract_text(in_path).strip()
    if len(cv_text) < 40:
        sys.exit("::error::Could not extract meaningful text from the CV.")
    print(f"       Extracted {len(cv_text)} characters.")

    print("[2/3] Calling Gemini …")
    enhanced = enhance_with_gemini(cv_text, args.prompt)
    print(f"       Received {len(enhanced)} characters.")

    stamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")
    stem = in_path.stem
    out_md = out_dir / f"{stem}_enhanced_{stamp}.md"

    header = (
        f"<!-- Enhanced CV\n"
        f"     Source : {in_path.name}\n"
        f"     For    : {args.email or 'n/a'}\n"
        f"     Prompt : {args.prompt}\n"
        f"     Model  : {os.environ.get('GEMINI_MODEL', 'gemini-2.0-flash')}\n"
        f"     Time   : {stamp} UTC\n"
        f"-->\n\n"
    )
    out_md.write_text(header + enhanced, encoding="utf-8")

    print(f"[3/3] Wrote enhanced CV -> {out_md}")

    # Expose the output path to later workflow steps if running in Actions
    gh_out = os.environ.get("GITHUB_OUTPUT")
    if gh_out:
        with open(gh_out, "a", encoding="utf-8") as fh:
            fh.write(f"enhanced_path={out_md}\n")


if __name__ == "__main__":
    main()
